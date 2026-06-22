#!/usr/bin/env python3
"""
Generate text-based test files:
- txt/sample.txt — plain text from sources/sample-source.md (stripped of markdown)
- pdf/sample.pdf — multi-page PDF via reportlab (headers, paragraphs, table)
- pdf/sample-scanned.pdf — image-only PDF (no text layer) via reportlab + PIL image embed
- docx/sample.docx — DOCX with headings + paragraphs + table via python-docx

Dependencies (install if missing):
  pip install reportlab python-docx pillow
"""
import re
import sys
from pathlib import Path

# Output base dir: first CLI arg (e.g. the target plan's test-data/), else current dir.
# Optional second arg: path to a markdown source file (else a built-in sample is used).
ROOT = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else Path.cwd()
SOURCE = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else ROOT / "sources" / "sample-source.md"
TXT_DIR = ROOT / "txt"
PDF_DIR = ROOT / "pdf"
DOCX_DIR = ROOT / "docx"

FALLBACK_SOURCE = """# Sample Document

This is a sample test document generated for Citeck feature acceptance testing.

## Section 1: Overview

The quick brown fox jumps over the lazy dog. This paragraph contains enough text to
exercise text extraction (pdftotext, DOCX parsing) and multimodal analysis flows.

## Section 2: Details

- First bullet point with some content.
- Second bullet point referencing **bold** and *italic* text.
- Third bullet point with a `code span`.

## Section 3: Table-like content

Item A - value 1
Item B - value 2
Item C - value 3
"""

for d in (TXT_DIR, PDF_DIR, DOCX_DIR):
    d.mkdir(parents=True, exist_ok=True)


def read_source():
    if SOURCE.exists():
        return SOURCE.read_text(encoding="utf-8")
    print(f"  (source {SOURCE} not found — using built-in sample text)")
    return FALLBACK_SOURCE


def make_txt(md: str):
    # Strip markdown: headers, code fences, inline backticks, links
    text = re.sub(r"^#{1,6}\s*", "", md, flags=re.MULTILINE)
    text = re.sub(r"```[^\n]*\n.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\|", " ", text)  # strip table pipes
    out = TXT_DIR / "sample.txt"
    out.write_text(text.strip() + "\n", encoding="utf-8")
    print(f"✓ {out} ({len(text.split())} words)")


def _register_unicode_font():
    """Register a TTF that supports Cyrillic. Returns the font family name."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    candidates = [
        # macOS
        ("/System/Library/Fonts/Supplemental/Arial.ttf",
         "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
         "/System/Library/Fonts/Supplemental/Arial Italic.ttf"),
        # Linux (Debian/Ubuntu)
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
        # Linux (other distros)
        ("/usr/share/fonts/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",
         "/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf"),
    ]
    import os
    for normal, bold, italic in candidates:
        if os.path.exists(normal):
            family = "UnicodeFont"
            pdfmetrics.registerFont(TTFont(family, normal))
            if os.path.exists(bold):
                pdfmetrics.registerFont(TTFont(f"{family}-Bold", bold))
            if os.path.exists(italic):
                pdfmetrics.registerFont(TTFont(f"{family}-Italic", italic))
            registerFontFamily(
                family, normal=family,
                bold=f"{family}-Bold" if os.path.exists(bold) else family,
                italic=f"{family}-Italic" if os.path.exists(italic) else family,
            )
            return family
    raise RuntimeError(
        "No Unicode font found. Install ttf-dejavu (Linux) or rely on macOS Arial."
    )


def make_pdf(md: str):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        )
        from reportlab.lib import colors
    except ImportError:
        print("⚠ reportlab not installed — pip install reportlab")
        return
    font_family = _register_unicode_font()
    out = PDF_DIR / "sample.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4)
    styles = getSampleStyleSheet()
    # Override default fonts for Cyrillic support
    for style_name in ("Title", "Heading1", "Heading2", "Heading3", "BodyText", "Normal"):
        if style_name in styles.byName:
            styles[style_name].fontName = font_family
    story = []
    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles["Title"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles["Heading2"]))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles["Heading3"]))
        elif line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            t = Table([cells], hAlign="LEFT")
            t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), font_family),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            story.append(t)
        else:
            story.append(Paragraph(line, styles["BodyText"]))
    doc.build(story)
    print(f"✓ {out}")


def make_scanned_pdf():
    """Image-only PDF — no text layer (для F-PM2 OpenAI text-fallback edge-case)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Image as RLImage, PageBreak
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("⚠ reportlab/pillow not installed")
        return
    out = PDF_DIR / "sample-scanned.pdf"
    pages = []
    for i in range(2):
        page_img = Image.new("RGB", (1240, 1754), "white")  # ~A4 200dpi
        draw = ImageDraw.Draw(page_img)
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 48)
        except OSError:
            font = ImageFont.load_default()
        draw.text((100, 100), f"Scanned page {i+1}", fill="black", font=font)
        for line_no, line in enumerate(["Citeck ECOS scanned page", "(text rendered as image, no text layer)"], 1):
            draw.text((100, 200 + line_no * 60), line, fill=(40, 40, 40), font=font)
        tmp = PDF_DIR / f"_scan_{i}.png"   # temp page image, unlinked after doc.build
        page_img.save(tmp)
        pages.append(tmp)
    # A4 inner frame is ~16cm × 24cm after default SimpleDocTemplate margins
    doc = SimpleDocTemplate(
        str(out), pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )
    story = []
    for i, p in enumerate(pages):
        story.append(RLImage(str(p), width=15 * cm, height=21 * cm))
        if i < len(pages) - 1:
            story.append(PageBreak())
    doc.build(story)
    for p in pages:
        p.unlink()
    print(f"✓ {out}")


def make_docx(md: str):
    try:
        from docx import Document
    except ImportError:
        print("⚠ python-docx not installed — pip install python-docx")
        return
    out = DOCX_DIR / "sample.docx"
    doc = Document()
    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            t = doc.add_table(rows=1, cols=len(cells))
            for c, val in zip(t.rows[0].cells, cells):
                c.text = val
        else:
            doc.add_paragraph(line)
    doc.save(str(out))
    print(f"✓ {out}")


if __name__ == "__main__":
    md = read_source()
    make_txt(md)
    make_pdf(md)
    make_scanned_pdf()
    make_docx(md)
    print("Done.")
